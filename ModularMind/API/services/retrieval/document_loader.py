"""
Belge yükleme işlemleri.
"""

import os
import uuid
from typing import List, Dict, Any, Optional
from fastapi import UploadFile
import time

from ModularMind.API.services.retrieval.models import Document

async def load_document_from_file(file: UploadFile) -> Document:
    """
    Yüklenen dosyadan belge oluşturur.
    
    Args:
        file: Yüklenen dosya
        
    Returns:
        Document: Oluşturulan belge
    """
    # Dosya uzantısını al
    filename = file.filename
    file_extension = os.path.splitext(filename)[1].lower() if filename else ""
    
    # Dosya içeriğini oku
    content = await file.read()
    
    # Dosya uzantısına göre içeriği işle
    if file_extension in ['.txt', '.md', '.csv']:
        text = content.decode('utf-8')
        return create_text_document(text, filename, file_extension)
        
    elif file_extension in ['.pdf']:
        return await extract_pdf_content(content, filename)
        
    elif file_extension in ['.docx', '.doc']:
        return await extract_docx_content(content, filename)
        
    elif file_extension in ['.html', '.htm']:
        return extract_html_content(content.decode('utf-8'), filename)
        
    else:
        # Desteklenmeyen dosya tipi
        raise ValueError(f"Desteklenmeyen dosya tipi: {file_extension}")

def create_text_document(text: str, filename: str, extension: str) -> Document:
    """
    Metin içeriğinden belge oluşturur.
    """
    # Belge metadata'sı
    metadata = {
        "source": filename,
        "file_type": extension.lstrip('.'),
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "title": os.path.splitext(filename)[0]
    }
    
    # Belge ID'si
    doc_id = f"doc_{uuid.uuid4().hex}"
    
    # Belge oluştur
    document = Document(
        id=doc_id,
        text=text,
        metadata=metadata
    )
    
    return document

async def extract_pdf_content(content: bytes, filename: str) -> Document:
    """
    PDF dosyasından içerik çıkarır.
    """
    try:
        from pypdf import PdfReader
        from io import BytesIO
        
        # PDF dosyasını oku
        pdf_file = BytesIO(content)
        pdf_reader = PdfReader(pdf_file)
        
        # Metin içeriğini çıkar
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        
        # Metadata
        metadata = {
            "source": filename,
            "file_type": "pdf",
            "page_count": len(pdf_reader.pages),
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "title": os.path.splitext(filename)[0]
        }
        
        # Belge ID'si
        doc_id = f"pdf_{uuid.uuid4().hex}"
        
        # Belge oluştur
        document = Document(
            id=doc_id,
            text=text,
            metadata=metadata
        )
        
        return document
    except ImportError:
        raise ImportError("PDF işleme için pypdf kütüphanesi gereklidir. pip install pypdf")

async def extract_docx_content(content: bytes, filename: str) -> Document:
    """
    DOCX dosyasından içerik çıkarır.
    """
    try:
        from docx import Document as DocxDocument
        from io import BytesIO
        
        # DOCX dosyasını oku
        docx_file = BytesIO(content)
        doc = DocxDocument(docx_file)
        
        # Metin içeriğini çıkar
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Tablolardan metin çıkar
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
        
        # Metadata
        metadata = {
            "source": filename,
            "file_type": "docx",
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "title": os.path.splitext(filename)[0]
        }
        
        # Belge ID'si
        doc_id = f"docx_{uuid.uuid4().hex}"
        
        # Belge oluştur
        document = Document(
            id=doc_id,
            text=text,
            metadata=metadata
        )
        
        return document
    except ImportError:
        raise ImportError("DOCX işleme için python-docx kütüphanesi gereklidir. pip install python-docx")

def extract_html_content(html_content: str, filename: str) -> Document:
    """
    HTML içeriğinden metin çıkarır.
    """
    try:
        from bs4 import BeautifulSoup
        
        # HTML içeriğini parse et
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Başlık çıkar
        title = soup.title.string if soup.title else os.path.splitext(filename)[0]
        
        # Metin içeriğini çıkar (script ve style etiketlerini kaldır)
        for script in soup(["script", "style"]):
            script.extract()
        
        # Sadece metin içeriğini al
        text = soup.get_text()
        
        # Boşlukları düzenle
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        # Metadata
        metadata = {
            "source": filename,
            "file_type": "html",
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "title": title
        }
        
        # Belge ID'si
        doc_id = f"html_{uuid.uuid4().hex}"
        
        # Belge oluştur
        document = Document(
            id=doc_id,
            text=text,
            metadata=metadata
        )
        
        return document
    except ImportError:
        raise ImportError("HTML işleme için beautifulsoup4 kütüphanesi gereklidir. pip install beautifulsoup4")